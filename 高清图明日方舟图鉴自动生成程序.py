import docx
from docx import Document
from docx.shared import Inches
import urllib.request
import re
import os
from bs4 import BeautifulSoup
import requests
from PIL import Image
import datetime
import time


def open_url(url):
    req = urllib.request.Request(url)
    req.add_header('User-Agent', 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36')
    page = urllib.request.urlopen(req)
    html = page.read().decode('utf-8')
    return html


def get_img(html,doc1):
    #p = r'<img class="BDE_Image".*?src="([^"]*\.jpg)".*?>'
    #<img class="BDE_Image" src="https://imgsa.baidu.com/forum/w%3D580/sign=2fd1e20182d4b31cf03c94b3b7d7276f/07f454a98226cffc7b2b46f0b0014a90f703eaa5.jpg"
    #data-src="https://mmbiz.qpic.cn/mmbiz_png/8gM0JJibLqVQoplKsdch5ZEzSe95A5Q5tcBAXOkcqVBIEj6vWw0ia1EZHtk1D9PzTJknpz8ZTJ2wXG2J6QRmibRtA/640?wx_fmt=png" data-type="png"
    file = "./图片库2/"
    if not os.path.exists(file):
        os.mkdir(file)
    soup = BeautifulSoup(html, "html.parser")
    soup.prettify()
    #a = soup.find('div' , class_= 'swiper-swrapper')
    #print(a)
    mainget = soup.find('title')
    nameget = re.findall('<title>(.*?)_', str(mainget))
    print(str(nameget[0]))
    print('开始延迟')
    try:
        try:
            soup1 = BeautifulSoup(open_url('http://arknights.huijiwiki.com/wiki/' + str('us'+str(bytes((str(nameget[0])).encode('utf-8')))).upper().replace(r'\X', '%').replace('USB\'', '').replace('\'', '')), "html.parser")
            time.sleep(8)
            print('干员取图成功')
        except:
            soup1 = BeautifulSoup(open_url('http://arknights.huijiwiki.com/wiki/' + str('us'+str(bytes((str(nameget[0])).encode('utf-8')))).replace(r'\X', '%').replace('usb\'', '').replace('\'', '')), "html.parser")
            time.sleep(8)
            print('小车取图成功')
        print('http://arknights.huijiwiki.com/wiki/' + str('us'+str(bytes((str(nameget[0])).encode('utf-8')))).upper().replace(r'\X', '%').replace('USB\'', '').replace('\'', ''))
        soup1.prettify()
        mainget = soup1.find('div', class_='tab-content')
        imgget = re.findall('src="(.*?)"', str(mainget))
        print(imgget)
        s = 0
        print('s赋值成功')
        for eachimg in imgget:
            print(eachimg)
            s += 1
            print(s)
            if s== 3:
                print('s==3'+'这是s:'+str(s))
                cddq = requests.get(eachimg)
                print('获取这张图片')
                with open(file + str(nameget[0]) + '精二' + ".png", 'wb')as f:
                    f.write(cddq.content)
                    f.close()
                try:
                    doc1.add_picture(file + str(nameget[0]) + '精二' + ".png", width=Inches(6.5))
                except:
                    doc1.add_paragraph('此图片不存在')
                    print('不存在')
            else:
                if s== 2:
                    print('s==2'+'这是s:'+str(s))
                    cddq = requests.get(eachimg)
                    print('获取这张图片')
                    with open(file + str(nameget) + '精一' + ".png", 'wb')as f:
                        f.write(cddq.content)
                        f.close()
                    try:
                        doc1.add_picture(file + str(nameget) + '精一' + ".png", width=Inches(6.5))
                    except:
                        doc1.add_paragraph('此图片不存在')
                        print(str('不存在'))
                else:
                    print('这张照片就算了')
    except:
            print('wait too long')
            doc1.add_paragraph('http://arknights.huijiwiki.com/wiki/' + str(bytes(str(nameget[0]).encode('utf-8'))).upper().replace(r'\X', '%').replace('B\'', '').replace('\'', ''))
            doc1.add_paragraph('此干员的图片没找到：'+str(nameget[0]))
    mainget = soup.find('table', class_="text")
    nextone = mainget.find_all('tr')
    print('nextone:' + str(nextone))
    titlecontent = re.findall('<tr><th>(.*?)</th>', str(nextone))
    nextone = mainget.find_all('td')
    doc1.add_paragraph('干员信息\n')
    cont = 0 - 1
    for eachone in titlecontent:
        cont += 1
        print(eachone + ":\n" + str(nextone[cont]).replace('<td colspan="5">', '').replace('</td>', ''))
        doc1.add_paragraph('\n')
        doc1.add_paragraph(eachone + ":\n" + str(nextone[cont]).replace('<td colspan="5">', '').replace('</td>', ''))
    mainget = soup.find('div', class_="char-right")
    nextone = mainget.find_all('tr')
    titlecontent = re.findall('<tr><th>(.*?)</th>', str(nextone))
    nextone = mainget.find_all('td')
    # doc1.add_paragraph('干员信息\n')
    cont = 0 - 1
    for eachone in titlecontent:
        cont += 1
        print(cont)
        print(eachone + ":\n" + str(nextone[cont]).replace('<td colspan="5"><a class="audio" href="javascript:;" rel="text_', '').replace('</td>', '').replace('"></a>', ''))
        doc1.add_paragraph('\n')
        doc1.add_paragraph(eachone + ":" + str(nextone[cont]).replace('<td colspan="5"><a class="audio" href="javascript:;" rel="text_', '').replace('</td>', '').replace('"></a>', ''))
    return doc1

if __name__ == '__main__':
    doc1 = docx.Document()
    for i in range(1,16):
        print('获取第'+ str(i) +'页')
        try:
            url1 = open_url("https://www.diopoo.com/ark/characters?pn=" + str(i))
            soup1 = BeautifulSoup(url1, "html.parser")
            soup1.prettify()
            list = soup1.find_all('div', class_='round shadow')
            charc = re.findall('href="(.*?)"',str(list))
            for eachone in charc:
                print('找到一个任务'+ str(eachone))
                url = "https://www.diopoo.com/ark/"+eachone
                doc1 = get_img(open_url(url),doc1)
        except:
            print('z此页无响应'+str(i))
    timeline = str(datetime.datetime.now().strftime('%Y.%m.%d'))
    doc1.save('./更新明日方舟全干员图鉴'+ timeline +'.docx')



