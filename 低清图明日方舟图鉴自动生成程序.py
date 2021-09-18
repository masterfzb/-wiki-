import docx
from docx import Document
from docx.shared import Inches
import urllib.request
import re
import os
from bs4 import BeautifulSoup
import requests
from PIL import Image
import datetime


def open_url(url):
    req = urllib.request.Request(url)
    req.add_header('User-Agent', 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36')
    page = urllib.request.urlopen(req)
    html = page.read().decode('utf-8')
    return html


def get_img(html,doc1,first,fake):
    #p = r'<img class="BDE_Image".*?src="([^"]*\.jpg)".*?>'
    #<img class="BDE_Image" src="https://imgsa.baidu.com/forum/w%3D580/sign=2fd1e20182d4b31cf03c94b3b7d7276f/07f454a98226cffc7b2b46f0b0014a90f703eaa5.jpg"
    #data-src="https://mmbiz.qpic.cn/mmbiz_png/8gM0JJibLqVQoplKsdch5ZEzSe95A5Q5tcBAXOkcqVBIEj6vWw0ia1EZHtk1D9PzTJknpz8ZTJ2wXG2J6QRmibRtA/640?wx_fmt=png" data-type="png"
    file = "./图片库/"
    if not os.path.exists(file):
        os.mkdir(file)
    soup = BeautifulSoup(html, "html.parser")
    soup.prettify()
    #a = soup.find('div' , class_= 'swiper-swrapper')
    #print(a)
    p = soup.find_all('div' , class_= 'swiper-slide')
    imglist = re.findall('"background-image:url\((.*?)\);"', str(p))
    print(imglist)
    i=0
    for eachless in imglist:
        i+=1
        if i%3 == 0 :
            print('这张照片不要')
        else:
            each = 'https://www.diopoo.com/ark/'+ eachless
            print("图片下载"+each)
            cddq = requests.get(each)
            with open(file + str(i + first + fake - 2) + ".png", 'wb')as f:
                f.write(cddq.content)
                f.close()
            try:
                doc1.add_picture(file + str(i + first + fake - 2) + ".png", width=Inches(6.5))
            except:
                doc1.add_paragraph('此图片不存在')
                print(str(i + first + fake - 2)+'不存在')
    mainget = soup.find('table', class_="text")
    print(mainget)
    nextone = mainget.find_all('tr')
    print('nextone:' + str(nextone))
    titlecontent = re.findall('<tr><th>(.*?)</th>', str(nextone))
    print(titlecontent)
    nextone = mainget.find_all('td')
    print(nextone[0])
    doc1.add_paragraph('干员信息\n')
    cont = 0 - 1
    for eachone in titlecontent:
        cont += 1
        print(cont)
        print(eachone + ":\n" + str(nextone[cont]).replace('<td colspan="5">', '').replace('</td>', ''))
        doc1.add_paragraph('\n')
        doc1.add_paragraph(eachone + ":\n" + str(nextone[cont]).replace('<td colspan="5">', '').replace('</td>', ''))
    mainget = soup.find('div', class_="char-right")
    print(mainget)
    nextone = mainget.find_all('tr')
    print('nextone:' + str(nextone))
    titlecontent = re.findall('<tr><th>(.*?)</th>', str(nextone))
    print(titlecontent)
    nextone = mainget.find_all('td')
    print(nextone[0])
    # doc1.add_paragraph('干员信息\n')
    cont = 0 - 1
    for eachone in titlecontent:
        cont += 1
        print(cont)
        print(eachone + ":\n" + str(nextone[cont]).replace(
            '<td colspan="5"><a class="audio" href="javascript:;" rel="text_', '').replace('</td>', '').replace(
            '"></a>', ''))
        doc1.add_paragraph('\n')
        doc1.add_paragraph(eachone + ":" + str(nextone[cont]).replace(
            '<td colspan="5"><a class="audio" href="javascript:;" rel="text_', '').replace('</td>', '').replace(
            '"></a>', ''))

if __name__ == '__main__':
    doc1 = docx.Document()
    for i in range(1,20):
        try:
            deit =(i-1)*32
            url1 = open_url("https://www.diopoo.com/ark/characters?pn=" + str(i))
            soup1 = BeautifulSoup(url1, "html.parser")
            soup1.prettify()
            list = soup1.find_all('div', class_='round shadow')
            print(list)
            charc = re.findall('href="(.*?)"',str(list))
            print(charc)
            fake = 0
            for eachone in charc:
                fake = fake + 2
                url = "https://www.diopoo.com/ark/"+eachone
                get_img(open_url(url),doc1,deit,fake)
        except:
            print(i)
    timeline = str(datetime.datetime.now().strftime('%Y.%m.%d'))
    doc1.save('./明日方舟全干员图鉴'+ timeline +'.docx')



